"""
        That file contains generic coupling modules:

         D -> Darcy
         R -> Richards
        2P -> 2 phases
         C -> Chemistry
         M -> Mechanics

"""
#
#
#
from __future__ import absolute_import
from __future__ import print_function
from six.moves import input
try: 
    from docx import Document
    from docx.shared import Inches,Cm
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.section import WD_ORIENTATION
except:
    raise Warning("You won't be able to handle report documents")
#
# The Python Imaging Library: This script can be used to change one image to another or remove an image.
#
try:
    import Image
except:
    from PIL import Image
#
# Return a list of paths matching a pathname pattern
#
from glob import glob

from listtools import toList

from os import getcwd, listdir, path, remove

from platform import architecture, uname

from types import ListType,\
                  MethodType,\
                  ModuleType

import resource

try:
    import shutil
except ImportError:
    # python version should be 2.7 or newer ones.
    pass

import subprocess

from sys import version_info

global pyversion

global plotId

pyversion = version_info[0] + version_info[1]*0.1

plotId = 0

class Color:
    """
    to print colors; example : print color.bold+"what you want to print" + color.end
    """
    purple = '\033[95m'
    cyan = '\033[96m'
    darkcyan = '\033[36m'
    blue = '\033[94m'
    green = '\033[92m'
    yellow = '\033[93m'
    red = '\033[91m'
    bold = '\033[1m'
    underline = '\033[4m'
    end = '\033[0m'

color = Color()

class Dico(object):
    """ Dico class definition as dictionary, keys are fifo organised (First In, First Out)
    """
    #---------------------------------------------------------------------------
    def __init__(self, name=''):
        """
        Constructor
        """
        self._name = name
        self._dict = {}
        self._cles = []
        pass
    #---------------------------------------------------------------------------
    def getName(self):
        """
        dictionary name, see the class constructor
        """
        return self._name
    #---------------------------------------------------------------------------        

    def keys(self):
        """
        list of boundary keys
        """
        return self._cles
    #---------------------------------------------------------------------------
    def values(self):
        """
        List of associated key values
        """
        listValues = []
        for  key in self._cles : listValues.append(self._dict[key])
        return listValues
    #---------------------------------------------------------------------------    
    def items(self):
        """ delivers a list of tuples:(key, value), key an entry of
            dico and value the associated valuee.
        """             
        listItems = []
        for  key in self._cles :listItems.append((key, self._dict[key]))
        return listItems
    #---------------------------------------------------------------------------
    def has_key(self, key):
        """ boolean 1 ifi key is a key of dico and 0 otherwise.
        """
        return key in self._dict
    #---------------------------------------------------------------------------
    def __len__(self):
        """ number of entries in the dico
        """
        return len(self._cles)
    #---------------------------------------------------------------------------
    def __setitem__(self, key, value):
        """ affectation of 'value' to the key 'key' within dico
            via the [] operator.
        """
        if key not in self._dict: self._cles.append(key)
        self._dict[key] = value
    #---------------------------------------------------------------------------

    def __getitem__(self, key):
        """
        return the value associated to 'key' via the [] operator.
        """
        if key in self._dict:
            return self._dict[key]
        else:
            return None
    #---------------------------------------------------------------------------
    def __str__(self):
        """
        Instance conversion:
        from an instance of the Dico class, we make a string.
        """
        lst = []
        for key  in self._cles:
            lst.append("%s:%s" % (key, self._dict[key]))
        ret = "Dico<name=%s>{%s}" % (self._name, string.join(lst, ", "))
        return ret
    #---------------------------------------------------------------------------   
    def __repr__(self):
        """
        """
        return  self.__str__()

    def delete(self,key):
        """
        We delete on element of the dictionary
        """
        del self._dict[key]
        self._cles.remove(key)
        return  
    
class Generic(object):
    """
    That class is only used as a generic container
    to provide help on class and associated methods to the user
    """
    def __init__(self):
        #self.pyversion = version_info[0] + version_info[1]*0.1
        #global pyversion
        self.pyversion = pyversion
        pass

    def getHelp(self, method = None):
        """
        That function enables to get some help on the
        class and on relevant methods:        
            a = class()
            getHelp(a.method)        
        Ex: getHelp() or getHelp(a.function)
        """
        if method == None or type(method) != MethodType:
            print(self.__doc__)
        else:
            print(method.__doc__)
        pass
        
    def clearVtkResults(self):
        """
        used to clear vtk files entailed in the current directory
        """
        name = getcwd()
        if path.exists(path.join(name,"VTK")):
            self.rmTree(path.join(name,"VTK"))
            pass
        vtkFiles = [fichier for fichier in listdir(os.getcwd()) if path.isfile(path.join(os.getcwd(), fichier))]
        removeFile(vtkFile for vtkFile in vtkFiles)
            
    def getModule(self):
        """
        To obtain the module name, the class is coming from
        """
        return self.__module__
        
    def rawInput(self,arg):
        return input("dbg"+self.__class__.__name__ +" "+str(arg))

    def removeFile(self, targetedFile):
        """
        Remove a file if it exists.
        """
        if path.exists(targetedFile):
            remove(targetedFile)
            pass

    def rmTree(self, targetDirectory, ignoreErrors = False):
        """
        Used to remove files in the directory tree.
        It wraps the shutil rmtree
        """
        if path.isdir(targetDirectory):
            shutil.rmtree(targetDirectory, ignore_errors = ignoreErrors)
        else:
            self.removeFile(targetDirectory)                                                         # removing the file targetDirectory
    
class GenericModule(Generic):
    """
    That class is only used as a generic container for new modules
    
    Modules are supposed to be transient by default 
    """
    def __init__(self) :
        """
        
        """
        Generic.__init__(self)
        
        self.spatialInteractiveOutputs = None 
        self.spatialSaveOutputs = None  
        self.vtkFileWriter = None
        self.vtkFrequency = None
    
        self.outputs = {}
        
        self.maxTimeStep = 1.e+15
        self.minTimeStep = 1.e-15
        
        self.simulatedTime = None
        self.timeStepNumber = None
        #
        self.debug = 0
        
    def cpuTime(self):
        """
        Used to determine the cpu time
        """
        return resource.getrusage(resource.RUSAGE_SELF)[0]
        
#    def clearResults(self,

class GenericCTModule(GenericModule):
    """
    That class is only used as a generic container
    for new modules
    """
    def __init__(self) :
        """
        
        """
        GenericModule.__init__(self)
        self.dT = None
        self.dispersivityBody = None      
        self.effectiveDiffusionZone = None
    
        self.vtkFieldSpecies = None
        
        self.chemicalOutputs = []      
        self.chemicalOutputNames = {}
        self.chemicalOutputDict = {}

class GenericWD_CTModule(GenericModule):
    """
    That class is only used as a generic container
    enabling a transient one phase Darcy coupled to a chemical transport module.
    """
    def __init__(self) :
        """
        
        """
        GenericModule.__init__(self)

def currentDirectory():
    return os.getcwd()

    
def checkClass(x, classes, message=None):
    """
    to check if x is not an instance of one of the classes list 
    """
    try:
        for klasse in classes:
            if x.__class__ == klasse.__class__:
                return
        pass
    except:
        pass
    if not isInstance(x, classes):
        if message is None:
            if hasattr(x, "__repr__"): xstr = x.__repr__()
            else: xstr = repr(type(x))
            xstr = repr(type(x))       
        raise TypeError(" x of type "+xstr+" is not within ["+_classesNameVerbose(classes)+"]")
    return
    
def reverseList(listToReverse):
    """
    cpu time efficiency > list(reversed(listToReverse)) and listToReverse[::-1]
    tested with the timeit module:
    
    timeit(lambda: reverseList(listToReverse), number=1000)
    
    """
    listToReverse.reverse()
    return listToReverse

def checkClassList(liste, classes):
    """
    to check if elements of liste  are instances of classes list classes
    """
    liste = list(liste)
    for x in liste:
        checkClass(x, classes)
    return None
    
def checkList(x, liste):
    """
    x should be an element of the list argument: liste
    Raises an exception if x is not in liste.
    """
    if x not in liste :
        l = liste.__str__()
        raise Exception(x + " not in list argument "+ l)
    pass
    
listCheck = checkList
    
def checkDict (x, dictionnary):
    """Raises an exception if an item is not a dictionnary key."""
    if x not in list(dictionary.keys()) :
        raise "key " + repr(x) + " not included in given dictionary"

dictCheck = checkDict
    
def fileAOpening(fileName):

    try:
        meshFile = open(fileName,"a")
    except IOError: " file opening error "
    return meshFile
        
def fileROpening(fileName):

    try:
        meshFile = open(fileName,"r")
    except IOError: " file opening for readind error "
    return meshFile

def fileWOpening(fileName):

    try:
        meshFile = open(fileName,"w")
        meshFile.write("##")
    except IOError:
        " file opening error "
    print(type(meshFile))
    return meshFile

def isInstance(x, klassenprobe):
    """Returns true if x is an instance of one of the classes within the testclasses, false otherwise"""
    if type(klassenprobe) is not ListType:
        if isinstance(x, klassenprobe):
            return 1
    else:
        for klasse in klassenprobe:
            if isinstance(x, klasse):
                return 1
    return 0
    
def kafka():
    """
    Enables to retrieve some elements of the Kafka bibliography
    a snapshot of the knowledge of the creator of the platform 
    """
    print("the process\nthe castle\nthe metamorphosis\n")
    return None

def listTypeCheck(liste, typ):
    if type(liste) != ListType:
        raise TypeError(" the supposed list is not of type list")
    if type(typ) != ListType:
        typ = [typ]
    for unb in liste:
        if type(unb) not in typ:
            print(type(unb),typ)
            raise TypeError(" the listed components are not all of the right type")
        pass
    return None
    
def memberShip(x,liste,message = None):
    if not isinstance(x,tuple(toList(liste))):
        if message == None: message = "of "+x.__class__.__name__
        raise Exception("check instantiation %s %s"%(message, x))
        
def makeDict(**options):
    """
    Example: makeDict(a=1,b=2) -> {'a': 1, 'b': 2}
    """
    return options

makeDico = makeDict

def postPlotter(dataTobePlotted, plotTitle = None, subTitle = None, time = None, gnuplotVersion = None, screen = None):
    """
    That function is used to save results in a png format.
    It is linked to the automatic insertion of simulation results in a doc like formatted file.
    gnuplot before the 5th version doesn't really support png format, even though those versions claim it.
    It means the plot generated is transformed in a png file through the image module.
    
    the entry should be a list od tuples (abs, ord)
    """
    global plotId, pyversion
    if gnuplotVersion == None:
        outputFormat = "eps"
    elif int(gnuplotVersion[0]) > 5:
       outputFormat = "png"
    fileName = "toto"
    gnuplot = open(fileName, "w")
    gnuplot.write("reset\n")
    gnuplot.write("set terminal %s\n"%(outputFormat))
    gnuplot.write("set style data linespoints\n")
    output = "output"
    if time != None:
        output = output+"_"+str("%i"%(plotId))+".png"
        plotId+=1
    gnuplot.write("set output \"%s\"\n"%(output))
    if plotTitle != None:
        gnuplot.write("set title \"" + str(plotTitle) + " (time = + %9.2e + years)\\n"%(time/3.15576e+7))
        if subTitle != None:
            gnuplot.write(str(subTitle)+"\" font \"Times-Roman,10\"\n")
        else:
            gnuplot.write("\" font \"Times-Roman,10\"\n")
    gnuplot.write("plot '-' using 1:2 title \"%s\" with linespoints \n"%(plotTitle))
    for point in dataTobePlotted:
        gnuplot.write("%f %f\n" % (point[0],point[1]))
    gnuplot.write("e\n")
    if screen != None:
        gnuplot.write("unset output\n")
        gnuplot.write("set terminal  dumb 79 25\n")
        gnuplot.write("replot\n")
    gnuplot.write("exit\n")
    gnuplot.flush()
    subprocess.Popen("gnuplot " + str(fileName), shell = True).wait()
    #gnuplot.close()
    img=Image.open(output)
    img.save(output)
    gnuplot.close()
    return output

def reportAddendum(reportFile,listofOutputsTobeAdded,listOfComments = None):
    """
    Enables to  add figures to a report in an automated way.
    The changes intervine after the last chapter of the report.
    """
    if ".docx" in reportFile:
        document = Document(reportFile)
    else:
        raise Warning(" problem with the format of the file")
    paragraph_styles = [s for s in document.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]

    postProcessingHeader = document.add_paragraph("PostProcessing", style=paragraph_styles[1])
    postProcessingText = document.add_paragraph(style=paragraph_styles[0])
    postProcessingText.add_run("\rWord format postprocessing generated through the ")
    postProcessingText.add_run("etumos ").bold = True
    postProcessingText.add_run("coupling tool\r\r")
    #document.add_heading('Outputs', level=1)
    ind = 0
    for picture in listofOutputsTobeAdded:
        document.add_picture(picture, width=Cm(7.00))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        newParagraph = document.add_paragraph("\r")
        newParagraph.add_run("Figure n. "+str(ind)+":   ").bold = True
        if listOfComments!= None:
            newParagraph.add_run(str(listOfComments[ind]))
        ind+=1
        newParagraph.add_run(" \r \r")
        newParagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    reportFile = reportFile.replace(".docx","")
    filenameToBeSaved = reportFile+"_pp.docx"
    document.save(filenameToBeSaved)
    return None
    
def keyAnalysis(ToolParameterAnalysisDictionary, key, value):
    """
    used to retrieve information from the user keywords
    to set solver or scheme parameters
    An example:
    
     "pressureSolverTolerance = value" becomes a tuple ("solver", "pressure", "Tolerance", value)
     returned by this method.
    """
    dico = {};dicoKey = {}
    if key:
        return dico, dicoKey, value
    else:
        return None, key, value
def _classesNameVerbose(liste):
    """returns a string made of the different classes names within a list,
       so it supposes each element has __name__ as attribute.
    """ 
    liste = toList(liste)
    return (' '.join([elt.__name__ for elt in liste]))

def IS_NUMBER(x):
    try:
        x = float(x)
        return True
    except:
        return False

def SET_NUMBER(x):
    if IS_NUMBER(x):
        return x
    else: raise Exception(" is not a number") 

def c_archi():
    if '32' in architecture()[0]:
        return '32'
    elif '64' in architecture()[0]:
        return '64'

def c_uname():
    """
        Returns a tuple
        of strings (system,node,release,version,machine,processor)
        identifying the underlying platform.
    """
    return uname

